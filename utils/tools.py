import os
import logging
from pathlib import Path

import docx
from docx.shared import Pt
from pydantic import BaseModel, Field

from langchain_core.tools import tool

from utils.dataframes import typo_data, subtypo_data, concept_data

MAIN_PATH = Path(os.getcwd())
DATA_PATH = MAIN_PATH / "data"
CASES_PATH = DATA_PATH / "cases"

# Logs
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(filename)s - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("app.log")
    ]
)
logger = logging.getLogger(__name__)



# Herramienta para saber si una tipologia se debe escalar o no -----------------------------------------------------------------
class GetTypoInfoInput(BaseModel):
    typo_code_list: int = Field(description="Code of the typology chosen")

@tool("get_typology_concept", args_schema=GetTypoInfoInput, return_direct=True)
def get_typology_concept(
        typo_code_list: int,
) -> str:

    """
    Returns if that tipology need a thrid-party concept
    and the casuistic, the area and additional requisites if any.

    Args:
        typo_code: Code of the typology chosen
    """
    logger.info("Tool get_third_part_concept used")

    logger.info(f"Typo code used: {typo_code_list}")
    concept_filtered = concept_data[concept_data["id"].isin([typo_code_list])].copy()
    # Si el df resulta vacio significa que esta tipologia no requiere concepto de tercero
    if concept_filtered.empty:
        typo_info = "La tipología buscada no fue encontrada"
        logger.error("Typology not found")
    # En caso contrario llevamos toda la info disponible para continuar
    else:
        typo_info = concept_filtered.copy().drop_duplicates()
        typo_info = "Estas son las tipologías encontradas:\n----\n" + "\n----\n".join(typo_info["data"].tolist())
        logger.info("Typology found")

    return typo_info


# Herramienta para saber obtener subtipologias -----------------------------------------------------------------
class GetSubtypoInfoInput(BaseModel):
    typo_code: int = Field(description="Code of the typology chosen")

@tool("get_subtypologies", args_schema=GetSubtypoInfoInput, return_direct=True)
def get_subtypologies(
        typo_code: int,
) -> str:

    """
    Returns all subtypologies the selectec typology has.

    Args:
        typo_code: Code of the typology chosen
    """
    logger.info("Tool get_subtypologies used")

    logger.info(f"Typo code used: {typo_code}")
    typo_filtered = subtypo_data[subtypo_data["id"].isin([typo_code])].copy()
    # Si el df resulta vacio significa que esta tipologia no requiere concepto de tercero
    if typo_filtered.empty:
        typo_info = "Esta tipología no tiene subtipologías"
        logger.error("Subtypologies not found")
    # En caso contrario llevamos toda la info disponible para continuar
    else:
        typo_info = typo_filtered.copy().drop_duplicates()
        typo_info = "Estas son las subtologías encontradas:\n----\n" + "\n----\n".join(typo_info["data"].tolist())
        logger.info("Subtypologies found")

    return typo_info


class MakeDocumentInput(BaseModel):
    date: str = Field(description="today's date")
    typo_name: str = Field(description="name of the typology chosen")
    typo_desc: str = Field(description="description of the typology chosen")
    pqrs_summary: str = Field(description="summary of the pqrs")
    file_name: str = Field(description="name of the pqrs document")

# Herramienta que recibe los datos encontrados por el agente 
# Y los convierte en la plantilla de respuesta del documento -----------------------------------------------------------------
@tool("make_response_document", args_schema=MakeDocumentInput, return_direct=True)
def make_response_document(
        date: str,
        typo_name: str,
        typo_desc: str,
        pqrs_summary: str,
        file_name: str
):
    """
    Returns the response document.

    Args:
        date: today's date
        typo_name: typology name
        typo_desc: typology description
        pqrs_summary: summary of the pqrs documents analized
        file_name: name of the pqrs document
    """
    logger.info("Tool make_response_document used")
     # 1. Crear un nuevo documento
    doc = docx.Document()

    try:
        # 2. Configurar el estilo de fuente por defecto (opcional, pero recomendado)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'  # Se asemeja a la fuente de la imagen
        font.size = Pt(11)

        # --- Encabezado ---
        doc.add_paragraph(f'Bogotá, ({date})')
        doc.add_paragraph('') # Espacio
        doc.add_paragraph('Señor(a)')
        doc.add_paragraph('Nombre del Cliente')
        doc.add_paragraph('Dirección del Cliente')
        doc.add_paragraph('Correo Electrónico')
        doc.add_paragraph('') # Espacio

        # --- Asunto ---
        asunto_p = doc.add_paragraph()
        asunto_p.add_run('Asunto: ').bold = False
        asunto_p.add_run(f'Respuesta a su reclamación sobre {typo_name}')

        doc.add_paragraph('') # Espacio

        # --- Cuerpo de la carta ---
        doc.add_paragraph('Estimado/a Nombre del Cliente,')
        doc.add_paragraph('Reciba un cordial saludo.')
        doc.add_paragraph('') # Espacio

        doc.add_paragraph(f'En atención a su solicitud relacionada con ({typo_desc})')
        doc.add_paragraph('') # Espacio

        # Detalle de la PQR (variable)
        doc.add_paragraph(pqrs_summary)
        doc.add_paragraph('') # Espacio

        # Párrafo con texto en negrita intercalado
        resolucion_p = doc.add_paragraph()
        resolucion_p.add_run('Sobre la resolución de su solicitud:')
        resolucion_p.add_run('\nConsiderando las circunstancias del caso, hemos decidido proceder con ')
        resolucion_p.add_run('decisión tomada: reversión de fondos, rechazo de la solicitud, etc').bold = True
        resolucion_p.add_run('. En consecuencia, se ha realizado ')
        resolucion_p.add_run('detalle de la acción tomada, por ejemplo, el abono en su cuenta por el valor reclamado o la negativa con sustento en la revisión del caso').bold = True
        resolucion_p.add_run('.')

        doc.add_paragraph('') # Espacio

        # --- Recomendaciones de seguridad ---
        doc.add_paragraph('Recomendaciones de seguridad:')
        recom_p = doc.add_paragraph('Para reforzar la seguridad de sus productos financieros, le sugerimos:')
        # Ajustar sangría para las viñetas
        recom_p.paragraph_format.left_indent = Pt(18)
        recom_p.paragraph_format.first_line_indent = Pt(-18)

        # Viñetas
        viñeta1 = doc.add_paragraph('- No compartir información confidencial como contraseñas o códigos de autenticación.')
        viñeta1.paragraph_format.left_indent = Pt(36)
        viñeta1.paragraph_format.first_line_indent = Pt(-18)

        viñeta2 = doc.add_paragraph('- Verificar la autenticidad de las llamadas y correos electrónicos de nuestra entidad.')
        viñeta2.paragraph_format.left_indent = Pt(36)
        viñeta2.paragraph_format.first_line_indent = Pt(-18)

        viñeta3 = doc.add_paragraph('- Revisar regularmente sus movimientos bancarios y reportar cualquier transacción sospechosa.')
        viñeta3.paragraph_format.left_indent = Pt(36)
        viñeta3.paragraph_format.first_line_indent = Pt(-18)

        doc.add_paragraph('Para el Banco su seguridad es importante, por eso le invitamos a ingresar al siguiente link https://bit.ly/2Vql6bt, donde encontrará todas las recomendaciones que en esta materia debe tener en cuenta para el uso de sus medios electrónicos.')
        doc.add_paragraph('') # Espacio

        # --- Canales de atención ---
        doc.add_paragraph('Canales de atención:')
        canales_p = doc.add_paragraph()
        canales_p.add_run('Si tiene alguna inquietud adicional, puede contactarnos a través de ')
        canales_p.add_run('canales de atención como teléfono, correo o aplicación móvil').bold = True
        canales_p.add_run('.')

        doc.add_paragraph('Recuerde que cuenta con nuestro canal transaccional BBVA Net al cual puede ingresar a través de www.bbva.com.co y nuestra APP BBVA Móvil, si necesita consultar o realizar transacciones de sus productos.')
        doc.add_paragraph('') # Espacio

        # --- Cierre ---
        doc.add_paragraph('Agradecemos su confianza en BBVA y reiteramos nuestro compromiso con la seguridad de sus productos financieros.')
        doc.add_paragraph('') # Espacio
        doc.add_paragraph('') # Espacio

        doc.add_paragraph('Atentamente,')
        doc.add_paragraph('') # Espacio para firma
        doc.add_paragraph('') # Espacio para firma
        doc.add_paragraph('') # Espacio para firma

        doc.add_paragraph('Nombre del Responsable').bold = True
        doc.add_paragraph('Departamento o Cargo')
        doc.add_paragraph('BBVA')
    except Exception as e:
        logger.error(f"Error while making response document: {e}")

    # 3. Guardar el documento
    case_path = CASES_PATH / file_name
    file_path = case_path / f"plantilla_respuesta_{file_name}.docx"
    doc.save(file_path)
    logger.info(f"Document created: {file_path.name}")

    return file_path